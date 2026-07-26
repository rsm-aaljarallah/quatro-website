import docx
import os

doc_path = os.path.abspath('AJ_AlJarallah_CV_Master.docx')
try:
    doc = docx.Document(doc_path)
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        if '2026' in p.text:
            p.text = p.text.replace('June 2026', 'December 2026')
            p.text = p.text.replace('Jun 2026', 'Dec 2026')
            p.text = p.text.replace('May 2026', 'December 2026')

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if '2026' in p.text:
                        p.text = p.text.replace('June 2026', 'December 2026')
                        p.text = p.text.replace('Jun 2026', 'Dec 2026')
                        p.text = p.text.replace('May 2026', 'December 2026')

    doc.save('AJ_AlJarallah_CV_Master.docx')
    print("Successfully updated docx")
except Exception as e:
    print(f"Error: {e}")
